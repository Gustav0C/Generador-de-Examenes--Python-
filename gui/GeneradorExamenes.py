import tkinter as tk
from tkinter import messagebox, simpledialog
import os
import string
from docx import Document
from docx2pdf import convert
import shutil
import re
import random
import platform
import subprocess
from PIL import Image, ImageTk
import customtkinter as ctk

class SoftwareExamenAdmision:
    def __init__(self, root):
        self.root = root
        self.root.title("v1.0")
        self.root.iconbitmap(os.path.join(os.path.dirname(__file__), "docs", "icon.ico"))
        self.root.geometry("600x450")
        self.root.resizable(False, False)

        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")

        # Frame principal
        self.frame = ctk.CTkFrame(self.root, corner_radius=15)
        self.frame.pack(padx=20, pady=20, fill="both", expand=True)

        # Título
        self.title_label = ctk.CTkLabel(self.frame, text="PANEL DE CONTROL", font=("Arial", 20, "bold"))
        self.title_label.pack(pady=15)

        # Imagen de título
        img_path = os.path.join(os.path.dirname(__file__), "docs", "1.png")
        if os.path.exists(img_path):
            img = Image.open(img_path)
            img = img.resize((200, 200), Image.Resampling.LANCZOS)
            img = ImageTk.PhotoImage(img)
            img_label = ctk.CTkLabel(self.frame, image=img, text="")
            img_label.image = img
            img_label.pack(pady=10) 

        # Botones con estilo mejorado
        self.ver_btn = ctk.CTkButton(self.frame, text="Ver Exámenes", command=self.ver_examenes)
        self.ver_btn.pack(pady=5, padx=30)

        self.generar_btn = ctk.CTkButton(self.frame, text="Generar Exámenes", command=self.generar_examenes)
        self.generar_btn.pack(pady=5, padx=30)

        self.calcular_btn = ctk.CTkButton(self.frame, text="Calcular Puntaje", command=self.calcular_puntaje_examenes)
        self.calcular_btn.pack(pady=5, padx=30)
        
        # Definir rutas
        self.docs_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs")
        self.examenes_generados_path = os.path.join(self.docs_path, "Examenes Generados")
        self.examen_original_path = os.path.join(self.docs_path, "Examen Original")
        self.plantilla_path = os.path.join(self.examen_original_path, "Examen Admision.docx")
        
        # Crear carpetas necesarias
        self.crear_carpetas()
        
    def calcular_puntaje_examenes(self):
        """Ejecutar el script externo de cálculo de puntaje."""
        try:
            subprocess.run(["python", r"C:\Users\Admin\dev\Generador de Examenes (Python)\gui\GeneradorDePuntajeFinal.py"], check=True)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo calcular el puntaje: {str(e)}")

    def generar_examenes(self):
        """Generar exámenes en PDF reemplazando el [TEMA] con letras desde A hasta Z."""
        try:
            # Verificar que exista la plantilla
            if not os.path.exists(self.plantilla_path):
                messagebox.showerror(
                    "Error", 
                    f"No se encontró la plantilla en {self.plantilla_path}. "
                    "Por favor, asegúrese de colocar la plantilla en esa ubicación."
                )
                return
            
            # Preguntar cantidad de exámenes a generar
            cantidad = simpledialog.askinteger(
                "Cantidad", 
                "¿Cuántos exámenes desea generar? (máximo 26 para temas A-Z)", 
                minvalue=1, 
                maxvalue=26
            )
            
            if not cantidad:
                return
            
            # Variable para controlar la cancelación
            self.cancelar_generacion = False

            #Guardar referencias a la ventana principal
            parent_window = self.root
            
            # Crear diálogo de progreso con mejor estilo
            progress_dialog = ctk.CTkToplevel(parent_window)
            progress_dialog.title("Generando Exámenes")
            progress_dialog.geometry("400x220")
            progress_dialog.transient(parent_window)
            
            # Centrar el diálogo en la pantalla
            x = parent_window.winfo_x() + (parent_window.winfo_width() - 400) // 2
            y = parent_window.winfo_y() + (parent_window.winfo_height() - 220) // 2
            progress_dialog.geometry(f"400x220+{x}+{y}")
            
            # Configurar el diálogo antes de grab_set
            progress_dialog.resizable(False, False)
            progress_dialog.protocol("WM_DELETE_WINDOW", lambda: None)

            # Frame principal para organizar elementos
            main_frame = ctk.CTkFrame(progress_dialog, corner_radius=10)
            main_frame.pack(fill="both", expand=True, padx=20, pady=20)

            # Título con ícono
            title_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
            title_frame.pack(fill="x", pady=(0, 10))
            
            title_label = ctk.CTkLabel(
                title_frame, 
                text="Generando Exámenes",
                font=ctk.CTkFont(family="Arial", size=16, weight="bold")
            )
            title_label.pack(side="left", padx=10)

            # Texto de progreso con mejor formato
            progress_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
            progress_frame.pack(fill="x", pady=5)
            
            status_icon = ctk.CTkLabel(
                progress_frame,
                text="🔄",  # Emoji como ícono
                font=ctk.CTkFont(size=18)
            )
            status_icon.pack(side="left", padx=(5, 0))

            # Texto de progreso
            progress_label = ctk.CTkLabel(
                progress_frame, 
                text="Preparando...", 
                font=ctk.CTkFont(family="Arial", size=14)
            )
            progress_label.pack(side="left", padx=10)

            # Contador de progreso
            counter_label = ctk.CTkLabel(
                main_frame, 
                text="0%",
                font=ctk.CTkFont(family="Arial", size=14, weight="bold")
            )
            counter_label.pack(pady=(0, 5))

            # Barra de progreso mejorada
            progress_bar = ctk.CTkProgressBar(
                main_frame, 
                width=350,
                height=15,
                corner_radius=5,
                mode="determinate"
            )
            progress_bar.pack(pady=10)
            progress_bar.set(0)  # Inicializar en 0%

            # Función para cancelar el proceso
            def cancelar_proceso():
                self.cancelar_generacion = True
                cancel_button.configure(text="Cancelando...", state="disabled")
                progress_label.configure(text="Cancelando operación...")
                status_icon.configure(text="⚠️")
                progress_dialog.update()

            # Botón de cancelar (desactivado durante el proceso)
            cancel_button = ctk.CTkButton(
                main_frame, 
                text="Cancelar",
                command=cancelar_proceso,
                width=100,
                corner_radius=8,
                fg_color="#c42b1c",
                hover_color="#971515"
            )
            cancel_button.pack(pady=10)

            # Asegurar que el diálogo está completamente dibujado
            for i in range(5):  # Múltiples actualizaciones forzadas
                progress_dialog.update()
                parent_window.update_idletasks()
                
            progress_dialog.grab_set()

            # Limpiar carpeta de exámenes temporales (si existe)
            temp_dir = os.path.join(self.docs_path, "temp")
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)
            
            # Generar los exámenes
            temas = list(string.ascii_uppercase)[:cantidad]  # A, B, C, ..., Z
            
            progress_label.configure(text="Iniciando generación...")
            progress_dialog.update()
            parent_window.update_idletasks()

            examenes_generados=0
            
            for i, tema in enumerate(temas):
                # Verificar si se ha cancelado el proceso
                if self.cancelar_generacion:
                    break

                # Actualizar barra de progreso
                porcentaje = int(((i + 1) / cantidad) * 100)
                progress_label.configure(text=f"Generando examen con tema {tema}...")
                counter_label.configure(text=f"{porcentaje}% ({i+1}/{cantidad})")
                
                if porcentaje >= 50:
                    status_icon.configure(text="⚡")  # Cambiar ícono al 50%
                
                progress_bar.set((i + 1) / cantidad)
                progress_dialog.update()
                
                # Crear una copia temporal de la plantilla
                temp_docx = os.path.join(temp_dir, f"Examen_Tema_{tema}.docx")
                
                # Cargar el documento
                doc = Document(self.plantilla_path)

                # Reemplazar tema en el documento actual
                for parrafo in doc.paragraphs:
                    for run in parrafo.runs:
                        texto = run.text
                        if texto and "[TEMA]" in texto:
                            nuevo_texto = texto.replace("[TEMA]", str(tema))
                            run.text = nuevo_texto

                # Guardar la versión modificada
                doc.save(temp_docx)

                # Convertir a PDF directamente
                pdf_output = os.path.join(self.examenes_generados_path, f"Examen_Tema_{tema}.pdf")
                convert(temp_docx, pdf_output)

                # Contar exámenes generados
                examenes_generados += 1

                progress_dialog.update()

            # Si se canceló el proceso
            if self.cancelar_generacion:
                status_icon.configure(text="⚠️")
                progress_label.configure(text="Operación cancelada")
                counter_label.configure(text=f"Generados: {examenes_generados}/{cantidad}")
                progress_dialog.update()
            else:
                # Cambiar ícono al completar
                status_icon.configure(text="✅")
                progress_label.configure(text="¡Proceso completado!")
                counter_label.configure(text="100%")
                progress_dialog.update()

            # Eliminar archivos temporales
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            
            # Cerrar diálogo de progreso
            progress_dialog.grab_release()
            progress_dialog.destroy()
            
            # Mostrar mensaje según resultado
            if self.cancelar_generacion:
                messagebox.showinfo("Operación Cancelada", f"Proceso cancelado. Se generaron {examenes_generados} exámenes.")
            else:
                # Mostrar diálogo de éxito personalizado en lugar de messagebox
                self.mostrar_dialogo_exito(f"Se han generado {cantidad} exámenes con temas del A al {temas[-1]} correctamente.")
                
        except Exception as e:
            if 'progress_dialog' in locals():
                progress_dialog.grab_release()
                progress_dialog.destroy()
            messagebox.showerror("Error", f"No se pudieron generar los exámenes: {str(e)}")

    def mostrar_dialogo_exito(self, mensaje):
        """Muestra un diálogo de éxito personalizado con CustomTkinter"""
        success_dialog = ctk.CTkToplevel(self.root)
        success_dialog.title("Éxito")
        success_dialog.geometry("400x220")
        success_dialog.transient(self.root)
        
        # Centrar el diálogo
        x = self.root.winfo_x() + (self.root.winfo_width() - 400) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 220) // 2
        success_dialog.geometry(f"400x220+{x}+{y}")
        
        # Configurar el diálogo
        success_dialog.resizable(False, False)
        
        # Frame principal
        main_frame = ctk.CTkFrame(success_dialog, corner_radius=10)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Icono de éxito
        success_icon = ctk.CTkLabel(
            main_frame,
            text="✅",
            font=ctk.CTkFont(size=36)
        )
        success_icon.pack(pady=(10, 5))
        
        # Mensaje de éxito
        message_label = ctk.CTkLabel(
            main_frame,
            text=mensaje,
            font=ctk.CTkFont(size=14),
            wraplength=340
        )
        message_label.pack(pady=10)
        
        # Botones
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill="x", pady=10)
        
        # Botón para ver exámenes
        view_button = ctk.CTkButton(
            button_frame,
            text="Ver Exámenes",
            command=lambda: [success_dialog.destroy(), self.ver_examenes()],
            width=120,
            corner_radius=8,
            hover_color="#2b5797"
        )
        view_button.pack(side="left", padx=10, expand=True)
        
        # Botón para cerrar
        close_button = ctk.CTkButton(
            button_frame,
            text="Cerrar",
            command=success_dialog.destroy,
            width=120,
            corner_radius=8,
            fg_color="#555555",
            hover_color="#333333"
        )
        close_button.pack(side="right", padx=10, expand=True)
        
        # Efecto de desvanecimiento
        success_dialog.attributes('-alpha', 0.0)
        success_dialog.update_idletasks()
        
        # Animar aparición
        for i in range(11):
            success_dialog.attributes('-alpha', i/10)
            success_dialog.update()

        # Iniciar con foco en el botón de ver exámenes
        view_button.focus_set()

    def crear_carpetas(self):
        """Crear las carpetas necesarias si no existen."""
        for path in [self.docs_path, self.examenes_generados_path, self.examen_original_path]:
            if not os.path.exists(path):
                os.makedirs(path)
        
        # Verificar si existe la plantilla original
        if not os.path.exists(self.plantilla_path):
            messagebox.showwarning(
                "Advertencia", 
                f"No se encontró la plantilla en {self.plantilla_path}. "
                "Por favor, asegúrese de colocar la plantilla en esa ubicación."
            )

    def extraer_preguntas_alternativas(self):
        # Cargar el documento
        doc = Document(self.plantilla_path)
        
        # Lista para almacenar preguntas y alternativas
        preguntas = []
        
        # Variables para seguimiento
        pregunta_actual = None
        alternativas = []
        
        # Patrones para identificar preguntas y alternativas según el formato específico
        patron_pregunta = re.compile(r'^(\d+\.)\s*(.*)', re.IGNORECASE)  # Formato: "1." o "2. "
        patron_alternativa = re.compile(r'^([a-e]\))\s*(.*)', re.IGNORECASE)  # Formato: "a)" hasta "e)"
        
        # Recorrer los párrafos del documento
        for parrafo in doc.paragraphs:
            texto = parrafo.text.strip()
            
            # Saltar párrafos vacíos
            if not texto:
                continue
            
            # Verificar si es una pregunta
            match_pregunta = patron_pregunta.match(texto)
            if match_pregunta:
                # Si hay una pregunta anterior, guardarla primero
                if pregunta_actual:
                    preguntas.append({
                        'pregunta': pregunta_actual,
                        'alternativas': alternativas.copy()
                    })
                
                # Iniciar nueva pregunta
                pregunta_actual = match_pregunta.group(2)
                alternativas = []
                continue
            
            # Verificar si es una alternativa
            match_alternativa = patron_alternativa.match(texto)
            if match_alternativa and pregunta_actual:
                letra = texto[0].upper()
                contenido = match_alternativa.group(2)
                alternativas.append({
                    'letra': letra,
                    'contenido': contenido
                })
        
        # Agregar la última pregunta si existe
        if pregunta_actual:
            preguntas.append({
                'pregunta': pregunta_actual,
                'alternativas': alternativas
            })
        
        return preguntas
    
    def reordenar_preguntas_alternativas(self, preguntas):
        # Copiar las preguntas para no modificar el original
        preguntas_reordenadas = preguntas.copy()
        
        # Reordenar aleatoriamente las preguntas
        random.shuffle(preguntas_reordenadas)
        
        # Para cada pregunta, reordenar sus alternativas
        for pregunta in preguntas_reordenadas:
            # Guardar las alternativas y reordenarlas
            alternativas = pregunta['alternativas'].copy()
            random.shuffle(alternativas)
            
            # Asignar nuevas letras según el orden
            letras = ['a', 'b', 'c', 'd', 'e']
            for i, alternativa in enumerate(alternativas):
                if i < len(letras):
                    alternativa['letra'] = letras[i]
            
            pregunta['alternativas'] = alternativas
        
        return preguntas_reordenadas
    
    def ver_examenes(self):
        """Abrir la carpeta de examenes generados."""
        try:
            if not os.path.exists(self.examenes_generados_path):
                os.makedirs(self.examenes_generados_path)
                
            # Abrir la carpeta según el sistema operativo
            if platform.system() == "Windows":
                os.startfile(self.examenes_generados_path)
            elif platform.system() == "Darwin":  # macOS
                subprocess.run(["open", self.examenes_generados_path])
            else:  # Linux
                subprocess.run(["xdg-open", self.examenes_generados_path])
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la carpeta: {str(e)}")